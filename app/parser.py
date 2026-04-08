import io
import os

import fitz  # PyMuPDF
import pytesseract
from PIL import Image
from docx import Document
from openpyxl import load_workbook

# Tesseractのパス設定
_tesseract_cmd = os.getenv("TESSERACT_CMD", r"D:\tools\Tesseract\tesseract.exe")
if os.path.exists(_tesseract_cmd):
    pytesseract.pytesseract.tesseract_cmd = _tesseract_cmd


SUPPORTED_EXTENSIONS = {".txt", ".pdf", ".docx", ".xlsx"}


def extract_text(filename: str, data: bytes) -> str:
    """ファイル形式に応じてテキストを抽出する"""
    ext = _get_ext(filename)
    if ext == ".pdf":
        return _extract_pdf(data)
    if ext == ".docx":
        return _extract_docx(data)
    if ext == ".xlsx":
        return _extract_xlsx(data)
    # デフォルト: プレーンテキスト
    return data.decode("utf-8")


def _get_ext(filename: str) -> str:
    return ("." + filename.rsplit(".", 1)[-1]).lower() if "." in filename else ""


def _extract_pdf(data: bytes) -> str:
    """PDFからテキストを抽出する（表はMarkdown形式で保持、テキストが少ない場合OCRにフォールバック）"""
    doc = fitz.open(stream=data, filetype="pdf")
    pages: list[str] = []
    for page_num, page in enumerate(doc, 1):
        parts: list[str] = [f"<!-- page:{page_num} -->"]
        text = page.get_text().strip()
        if len(text) < 20:
            text = _ocr_page(page)
        # 表を検出してMarkdown形式に変換
        tables_md = _extract_tables_from_page(page)
        if tables_md:
            parts.append(text)
            parts.extend(tables_md)
        else:
            parts.append(text)
        pages.append("\n\n".join(parts))
    doc.close()
    return "\n\n".join(pages)


def _extract_tables_from_page(page) -> list[str]:
    """PDFページから表を検出し、Markdown表形式のリストとして返す"""
    try:
        tables = page.find_tables()
        if not tables or len(tables.tables) == 0:
            return []
        result: list[str] = []
        for table in tables:
            rows = table.extract()
            if not rows or len(rows) < 2:
                continue
            md_rows: list[str] = []
            # ヘッダー行
            header = [str(c).strip() if c else "" for c in rows[0]]
            md_rows.append("| " + " | ".join(header) + " |")
            md_rows.append("| " + " | ".join(["---"] * len(header)) + " |")
            # データ行
            for row in rows[1:]:
                cells = [str(c).strip() if c else "" for c in row]
                # 列数をヘッダーに揃える
                while len(cells) < len(header):
                    cells.append("")
                md_rows.append("| " + " | ".join(cells[:len(header)]) + " |")
            result.append("\n".join(md_rows))
        return result
    except Exception:
        return []


def _ocr_page(page) -> str:
    """PDFページを画像化してOCRでテキスト抽出する"""
    try:
        pix = page.get_pixmap(dpi=300)
        img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
        text = pytesseract.image_to_string(img, lang="jpn+eng")
        return text.strip()
    except Exception:
        return ""


def _extract_docx(data: bytes) -> str:
    """Word(.docx)からテキストを抽出する（表はMarkdown形式で保持）"""
    doc = Document(io.BytesIO(data))
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    # 表を抽出してMarkdown形式に変換
    for table in doc.tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        if len(rows) < 2:
            continue
        md_rows: list[str] = []
        header = rows[0]
        md_rows.append("| " + " | ".join(header) + " |")
        md_rows.append("| " + " | ".join(["---"] * len(header)) + " |")
        for row in rows[1:]:
            cells = row
            while len(cells) < len(header):
                cells.append("")
            md_rows.append("| " + " | ".join(cells[:len(header)]) + " |")
        parts.append("\n".join(md_rows))
    return "\n\n".join(parts)


def _extract_xlsx(data: bytes) -> str:
    """Excel(.xlsx)からテキストを抽出する（シートごとに見出し付き）"""
    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    parts: list[str] = []
    for sheet in wb.worksheets:
        rows: list[str] = []
        for row in sheet.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            line = "\t".join(cells).strip()
            if line:
                rows.append(line)
        if rows:
            parts.append(f"## シート: {sheet.title}\n" + "\n".join(rows))
    wb.close()
    return "\n\n".join(parts)
