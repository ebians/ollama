"""parser.py のテスト（Ollama / Tesseract 不要）"""
import pytest
from app.parser import extract_text, _get_ext, _extract_docx, _extract_xlsx, SUPPORTED_EXTENSIONS


class TestGetExt:
    def test_txt(self):
        assert _get_ext("readme.txt") == ".txt"

    def test_pdf(self):
        assert _get_ext("document.PDF") == ".pdf"

    def test_docx(self):
        assert _get_ext("report.docx") == ".docx"

    def test_xlsx(self):
        assert _get_ext("data.xlsx") == ".xlsx"

    def test_no_extension(self):
        assert _get_ext("noext") == ""

    def test_multiple_dots(self):
        assert _get_ext("my.file.name.pdf") == ".pdf"


class TestSupportedExtensions:
    def test_contains_all(self):
        assert ".txt" in SUPPORTED_EXTENSIONS
        assert ".pdf" in SUPPORTED_EXTENSIONS
        assert ".docx" in SUPPORTED_EXTENSIONS
        assert ".xlsx" in SUPPORTED_EXTENSIONS

    def test_unsupported(self):
        assert ".exe" not in SUPPORTED_EXTENSIONS
        assert ".jpg" not in SUPPORTED_EXTENSIONS


class TestExtractText:
    def test_plain_text_utf8(self):
        content = "これはテストです\nHello World"
        result = extract_text("test.txt", content.encode("utf-8"))
        assert "これはテストです" in result
        assert "Hello World" in result

    def test_plain_text_ascii(self):
        content = "Simple ASCII text"
        result = extract_text("test.txt", content.encode("utf-8"))
        assert result == "Simple ASCII text"

    def test_empty_text(self):
        result = extract_text("empty.txt", b"")
        assert result == ""


class TestExtractPdf:
    def test_text_pdf(self):
        """テキストレイヤーのあるPDFからテキスト抽出"""
        import fitz
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((50, 50), "PDF test content", fontsize=12)
        pdf_bytes = doc.tobytes()
        doc.close()

        result = extract_text("test.pdf", pdf_bytes)
        assert "PDF test content" in result


class TestExtractDocx:
    def test_docx(self):
        """Word文書からテキスト抽出"""
        from docx import Document
        import io
        doc = Document()
        doc.add_paragraph("段落1のテキスト")
        doc.add_paragraph("段落2のテキスト")
        buf = io.BytesIO()
        doc.save(buf)

        result = extract_text("test.docx", buf.getvalue())
        assert "段落1のテキスト" in result
        assert "段落2のテキスト" in result


class TestExtractXlsx:
    def test_xlsx(self):
        """Excelファイルからテキスト抽出"""
        from openpyxl import Workbook
        import io
        wb = Workbook()
        ws = wb.active
        ws.title = "テストシート"
        ws["A1"] = "名前"
        ws["B1"] = "部署"
        ws["A2"] = "田中"
        ws["B2"] = "開発部"
        buf = io.BytesIO()
        wb.save(buf)

        result = extract_text("test.xlsx", buf.getvalue())
        assert "テストシート" in result
        assert "田中" in result
        assert "開発部" in result


class TestDocxTableExtraction:
    def test_docx_with_table(self):
        """Word文書の表がMarkdown形式で抽出される"""
        from docx import Document
        import io
        doc = Document()
        doc.add_paragraph("レポート概要")
        table = doc.add_table(rows=3, cols=2)
        table.cell(0, 0).text = "項目"
        table.cell(0, 1).text = "金額"
        table.cell(1, 0).text = "売上"
        table.cell(1, 1).text = "100万円"
        table.cell(2, 0).text = "利益"
        table.cell(2, 1).text = "20万円"
        buf = io.BytesIO()
        doc.save(buf)

        result = extract_text("report.docx", buf.getvalue())
        assert "レポート概要" in result
        assert "| 項目 | 金額 |" in result
        assert "| 売上 | 100万円 |" in result
        assert "| 利益 | 20万円 |" in result
        assert "| --- | --- |" in result


class TestPdfPageComment:
    def test_pdf_has_page_comment(self):
        """PDFテキスト抽出にページ番号コメントが含まれる"""
        import fitz
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((50, 50), "Page 1 content", fontsize=12)
        page2 = doc.new_page()
        page2.insert_text((50, 50), "Page 2 content", fontsize=12)
        pdf_bytes = doc.tobytes()
        doc.close()

        result = extract_text("test.pdf", pdf_bytes)
        assert "<!-- page:1 -->" in result
        assert "<!-- page:2 -->" in result
        assert "Page 1 content" in result
        assert "Page 2 content" in result
